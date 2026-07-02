"""讀取器單元測試:每種格式用真實產生的樣本檔驗證欄位與資料抽取正確。"""
from __future__ import annotations

import pytest

from app.parsers import ParseError, parse


def _write_xlsx(path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["廠商名稱", "統一編號", "金額"])
    ws.append(["大同營造", "12345678", 1000])
    ws.append(["中華工程", "87654321", 2000])
    wb.save(path)


def test_excel(tmp_path):
    p = tmp_path / "s.xlsx"
    _write_xlsx(p)
    r = parse(p)
    assert r.source_type == "excel"
    assert r.fields == ["廠商名稱", "統一編號", "金額"]
    assert len(r.rows) == 2
    assert r.rows[0]["廠商名稱"] == "大同營造"
    assert r.sample_value("統一編號") == "12345678"


def test_csv(tmp_path):
    p = tmp_path / "s.csv"
    p.write_text("廠商名稱,統一編號,金額\n大同營造,12345678,1000\n", encoding="utf-8-sig")
    r = parse(p)
    assert r.source_type == "csv"
    assert r.fields == ["廠商名稱", "統一編號", "金額"]
    assert r.rows[0]["金額"] == "1000"


def test_csv_big5(tmp_path):
    p = tmp_path / "s_big5.csv"
    p.write_bytes("廠商名稱,金額\n大同營造,1000\n".encode("big5"))
    r = parse(p)
    assert r.fields == ["廠商名稱", "金額"]
    assert r.rows[0]["廠商名稱"] == "大同營造"


def test_ods(tmp_path):
    pytest.importorskip("odf")
    from odf.opendocument import OpenDocumentSpreadsheet
    from odf.table import Table, TableRow, TableCell
    from odf.text import P

    doc = OpenDocumentSpreadsheet()
    table = Table(name="Sheet1")
    for row in [["廠商名稱", "金額"], ["大同營造", "1000"]]:
        tr = TableRow()
        for val in row:
            tc = TableCell(valuetype="string")
            tc.addElement(P(text=val))
            tr.addElement(tc)
        table.addElement(tr)
    doc.spreadsheet.addElement(table)
    p = tmp_path / "s.ods"
    doc.save(str(p))

    r = parse(p)
    assert r.source_type == "ods"
    assert "廠商名稱" in r.fields


def test_word_table(tmp_path):
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "廠商名稱"
    table.cell(0, 1).text = "金額"
    table.cell(1, 0).text = "大同營造"
    table.cell(1, 1).text = "1000"
    p = tmp_path / "s.docx"
    doc.save(str(p))

    r = parse(p)
    assert r.source_type == "word"
    assert r.fields == ["廠商名稱", "金額"]
    assert r.rows[0]["廠商名稱"] == "大同營造"


def test_word_keyvalue(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("廠商名稱:大同營造")
    doc.add_paragraph("統一編號:12345678")
    p = tmp_path / "kv.docx"
    doc.save(str(p))

    r = parse(p)
    assert set(r.fields) == {"廠商名稱", "統一編號"}
    assert r.rows[0]["統一編號"] == "12345678"


def test_unsupported_image(tmp_path):
    p = tmp_path / "scan.png"
    p.write_bytes(b"\x89PNG\r\n")
    with pytest.raises(ParseError) as exc:
        parse(p)
    assert "範圍外" in str(exc.value) or "不支援" in str(exc.value)


def test_missing_file(tmp_path):
    with pytest.raises(ParseError):
        parse(tmp_path / "nope.xlsx")

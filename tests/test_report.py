"""報表抽取與產出測試,含端對端流程(規格 §9)。"""
from __future__ import annotations

import pytest

from app import report
from app.parsers.base import ParseResult


# --- 欄位抽取 ------------------------------------------------------------

def test_extract_excel_placeholder(tmp_path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws["A1"] = "廠商:"
    ws["B1"] = "{{廠商名稱}}"
    ws["A2"] = "金額:"
    ws["B2"] = "{{金額}}"
    p = tmp_path / "tpl.xlsx"
    wb.save(p)

    fields = report.extract_fields(p, "excel")
    names = [f["field_name"] for f in fields]
    assert names == ["廠商名稱", "金額"]
    assert fields[0]["location"] == "B1"


def test_extract_excel_header(tmp_path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["廠商名稱", "統一編號", "金額"])
    p = tmp_path / "tpl_header.xlsx"
    wb.save(p)

    fields = report.extract_fields(p, "excel")
    assert [f["field_name"] for f in fields] == ["廠商名稱", "統一編號", "金額"]


def test_extract_word(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("廠商名稱:{{廠商名稱}}")
    doc.add_paragraph("金額:{{金額}}")
    p = tmp_path / "tpl.docx"
    doc.save(str(p))

    fields = report.extract_fields(p, "word")
    assert {f["field_name"] for f in fields} == {"廠商名稱", "金額"}


# --- 產出 ---------------------------------------------------------------

def _patch_output(monkeypatch, tmp_path):
    out = tmp_path / "output"
    out.mkdir()
    monkeypatch.setattr(report, "OUTPUT_DIR", out)
    return out


def test_generate_excel_placeholder(tmp_path, monkeypatch):
    from openpyxl import Workbook, load_workbook

    _patch_output(monkeypatch, tmp_path)
    wb = Workbook()
    ws = wb.active
    ws["B1"] = "{{廠商名稱}}"
    ws["B2"] = "{{金額}}"
    tpl = tmp_path / "tpl.xlsx"
    wb.save(tpl)

    report_fields = report.extract_fields(tpl, "excel")
    mappings = [
        {"report_field": "廠商名稱", "source_field": "名稱"},
        {"report_field": "金額", "source_field": "工程金額"},
    ]
    source_rows = [{"名稱": "大同營造", "工程金額": "9999"}]

    result = report.generate(tpl, "excel", report_fields, mappings, source_rows, "out.xlsx")
    assert result["status"] == "ok"
    wb2 = load_workbook(result["output_path"])
    ws2 = wb2.active
    assert ws2["B1"].value == "大同營造"
    assert ws2["B2"].value == "9999"


def test_generate_excel_header_multirow(tmp_path, monkeypatch):
    from openpyxl import Workbook, load_workbook

    _patch_output(monkeypatch, tmp_path)
    wb = Workbook()
    ws = wb.active
    ws.append(["廠商名稱", "金額"])
    tpl = tmp_path / "tpl.xlsx"
    wb.save(tpl)

    report_fields = report.extract_fields(tpl, "excel")
    mappings = [
        {"report_field": "廠商名稱", "source_field": "名稱"},
        {"report_field": "金額", "source_field": "金額"},
    ]
    source_rows = [
        {"名稱": "大同營造", "金額": "1000"},
        {"名稱": "中華工程", "金額": "2000"},
    ]
    result = report.generate(tpl, "excel", report_fields, mappings, source_rows, "out2.xlsx")
    wb2 = load_workbook(result["output_path"])
    ws2 = wb2.active
    assert ws2["A2"].value == "大同營造"
    assert ws2["A3"].value == "中華工程"
    assert ws2["B3"].value == "2000"


def test_generate_warns_on_unmapped(tmp_path, monkeypatch):
    from openpyxl import Workbook

    _patch_output(monkeypatch, tmp_path)
    wb = Workbook()
    ws = wb.active
    ws["B1"] = "{{完工日期}}"
    tpl = tmp_path / "tpl.xlsx"
    wb.save(tpl)

    report_fields = report.extract_fields(tpl, "excel")
    mappings = [{"report_field": "完工日期", "source_field": ""}]  # 未對應
    source_rows = [{"名稱": "大同營造"}]

    result = report.generate(tpl, "excel", report_fields, mappings, source_rows, "out3.xlsx")
    assert result["status"] == "warning"
    assert any("完工日期" in w for w in result["warnings"])

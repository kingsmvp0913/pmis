"""報表欄位抽取 + 依對應填入固定範本產出報表。

單一職責:
1. ``extract_fields`` —— 從報表固定範本抽出「這個報表需要哪些欄位」。
2. ``generate`` —— 依對應把來源資料填入範本,產出檔案(存 data/output/)。

支援兩種範本:
- **Excel**(openpyxl):
    * placeholder 模式:儲存格內含 ``{{欄位名}}`` → 該欄位;產出時以資料值取代。
    * 表頭模式:第一列為欄位名,資料由第二列起逐列填入。
- **Word**(docxtpl):以 ``{{欄位名}}`` jinja 佔位,產出時代入資料(取第一列)。

「範本欄位對不到來源資料 → 該欄留空並在產出記錄標註,不靜默略過」(見規格 §8)。
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from .base_paths import OUTPUT_DIR
from .parsers.base import ParseError

_PLACEHOLDER = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")


# --- 欄位抽取 ------------------------------------------------------------

def extract_fields(template_path: str | Path, template_type: str) -> list[dict[str, Any]]:
    """從報表範本抽出欄位清單:[{field_name, location, sort}]。"""
    path = Path(template_path)
    if not path.exists():
        raise ParseError(f"找不到報表範本:{path}")

    if template_type == "excel":
        return _extract_excel_fields(path)
    if template_type == "word":
        return _extract_word_fields(path)
    raise ParseError(f"不支援的報表範本類型:{template_type}")


def _extract_excel_fields(path: Path) -> list[dict[str, Any]]:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=False)
    ws = wb.active

    placeholder_fields: list[dict[str, Any]] = []
    seen: set[str] = set()
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value, str):
                for m in _PLACEHOLDER.finditer(cell.value):
                    name = m.group(1).strip()
                    if name and name not in seen:
                        seen.add(name)
                        placeholder_fields.append(
                            {
                                "field_name": name,
                                "location": cell.coordinate,
                                "sort": len(placeholder_fields),
                            }
                        )
    if placeholder_fields:
        wb.close()
        return placeholder_fields

    # 表頭模式:第一列即欄位
    header_fields: list[dict[str, Any]] = []
    first_row = next(ws.iter_rows(min_row=1, max_row=1), ())
    for cell in first_row:
        if cell.value not in (None, ""):
            header_fields.append(
                {
                    "field_name": str(cell.value).strip(),
                    "location": cell.coordinate,
                    "sort": len(header_fields),
                }
            )
    wb.close()
    if not header_fields:
        raise ParseError("Excel 報表範本中找不到 {{欄位}} 佔位,也沒有表頭欄位。")
    return header_fields


def _extract_word_fields(path: Path) -> list[dict[str, Any]]:
    from docx import Document

    doc = Document(str(path))
    seen: set[str] = set()
    fields: list[dict[str, Any]] = []

    def scan(text: str) -> None:
        for m in _PLACEHOLDER.finditer(text or ""):
            name = m.group(1).strip()
            if name and name not in seen:
                seen.add(name)
                fields.append(
                    {"field_name": name, "location": "docx", "sort": len(fields)}
                )

    for para in doc.paragraphs:
        scan(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan(cell.text)

    if not fields:
        raise ParseError("Word 報表範本中找不到 {{欄位}} 佔位。")
    return fields


# --- 產出 ---------------------------------------------------------------

def generate(
    template_path: str | Path,
    template_type: str,
    report_fields: list[dict[str, Any]],
    mappings: list[dict[str, Any]],
    source_rows: list[dict[str, Any]],
    output_filename: str,
    carry_over: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """依對應把來源資料填入範本產出檔案。

    mappings:[{report_field, source_field}] —— source_field 為空代表沒對到。
    source_rows:來源解析出的資料列。
    carry_over:上期同報表同廠商的欄位值;當本次某欄留空時,以此帶入(功能 5)。
    回傳:{output_path, warnings:[...], status, values}
      values:第一列填入的 {報表欄位: 值},供「上期帶入」儲存與稽核。
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / output_filename

    rf_to_sf = {m["report_field"]: (m.get("source_field") or "") for m in mappings}
    warnings: list[str] = []

    # 先算出每列的 {報表欄位: 值}(含上期帶入),再寫入範本 —— 邏輯與寫檔分離
    rows_values = _compute_rows_values(
        report_fields, rf_to_sf, source_rows, carry_over, warnings
    )

    if template_type == "excel":
        _fill_excel(template_path, out_path, report_fields, rows_values, warnings)
    elif template_type == "word":
        _fill_word(template_path, out_path, report_fields, rows_values, warnings)
    else:
        raise ParseError(f"不支援的報表範本類型:{template_type}")

    status = "warning" if warnings else "ok"
    return {
        "output_path": str(out_path),
        "warnings": warnings,
        "status": status,
        "values": rows_values[0] if rows_values else {},
    }


def _compute_rows_values(
    report_fields: list[dict[str, Any]],
    rf_to_sf: dict[str, str],
    source_rows: list[dict[str, Any]],
    carry_over: dict[str, Any] | None,
    warnings: list[str],
) -> list[dict[str, Any]]:
    """把每一列來源資料換算成 {報表欄位: 值};空值時嘗試以上期帶入。"""
    carry_over = carry_over or {}
    rows = source_rows if source_rows else [{}]
    result: list[dict[str, Any]] = []
    warned_missing: set[str] = set()

    for idx, row in enumerate(rows):
        values: dict[str, Any] = {}
        for f in report_fields:
            fn = f["field_name"]
            src = rf_to_sf.get(fn, "")
            val: Any = ""
            if not src:
                if fn not in warned_missing:
                    warnings.append(f"報表欄位「{fn}」沒有設定對應來源。")
                    warned_missing.add(fn)
            elif src not in row:
                if fn not in warned_missing:
                    warnings.append(
                        f"來源資料中找不到欄位「{src}」(對應報表欄位「{fn}」)。"
                    )
                    warned_missing.add(fn)
            else:
                val = row.get(src, "")

            # 上期帶入:僅第一列、且本次留空時
            if (val is None or val == "") and idx == 0 and carry_over.get(fn):
                val = carry_over[fn]
                warnings.append(f"報表欄位「{fn}」本次無資料,已帶入上期值。")
            values[fn] = val
        result.append(values)
    return result


def _fill_excel(
    template_path: str | Path,
    out_path: Path,
    report_fields: list[dict[str, Any]],
    rows_values: list[dict[str, Any]],
    warnings: list[str],
) -> None:
    from openpyxl import load_workbook
    from openpyxl.utils import coordinate_to_tuple

    wb = load_workbook(template_path)
    ws = wb.active

    # 判斷是 placeholder 模式還是表頭模式
    is_placeholder = False
    for f in report_fields:
        loc = f.get("location", "")
        if loc:
            try:
                cell = ws[loc]
            except Exception:
                continue
            if isinstance(cell.value, str) and _PLACEHOLDER.search(cell.value):
                is_placeholder = True
                break

    row0 = rows_values[0] if rows_values else {}

    if is_placeholder:
        for f in report_fields:
            loc = f["location"]
            field_name = f["field_name"]
            try:
                cell = ws[loc]
            except Exception:
                warnings.append(f"報表欄位「{field_name}」的位置 {loc} 無效,略過。")
                continue
            value = row0.get(field_name, "")
            if isinstance(cell.value, str):
                cell.value = _PLACEHOLDER.sub(
                    lambda m, fn=field_name, v=value: str(v) if m.group(1).strip() == fn else m.group(0),
                    cell.value,
                )
            else:
                cell.value = value
        if len(rows_values) > 1:
            warnings.append(
                f"來源有 {len(rows_values)} 列資料,placeholder 範本僅填入第一列。"
            )
    else:
        col_of: dict[str, int] = {}
        for f in report_fields:
            loc = f.get("location", "")
            try:
                _, c = coordinate_to_tuple(loc)
                col_of[f["field_name"]] = c
            except Exception:
                warnings.append(f"報表欄位「{f['field_name']}」位置 {loc} 無效,略過。")
        start_row = 2
        for i, values in enumerate(rows_values):
            for field_name, col in col_of.items():
                ws.cell(row=start_row + i, column=col, value=values.get(field_name, ""))

    wb.save(out_path)


def _fill_word(
    template_path: str | Path,
    out_path: Path,
    report_fields: list[dict[str, Any]],
    rows_values: list[dict[str, Any]],
    warnings: list[str],
) -> None:
    from docxtpl import DocxTemplate

    context = dict(rows_values[0]) if rows_values else {}
    if len(rows_values) > 1:
        warnings.append(f"來源有 {len(rows_values)} 列資料,Word 範本僅填入第一列。")

    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    tpl.save(str(out_path))

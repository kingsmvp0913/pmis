"""Word 文件讀取器(.docx)。

營造來源文件常見兩種排版,兩者都支援:
1. **表格**:第一列為表頭,其餘為資料列(與 Excel 類似)。
2. **鍵值段落**:如「廠商名稱:大同營造」「統一編號:12345678」,
   拆成 {欄位: 值} 的單列資料。

若同時存在,優先採用資料量較完整的表格;否則採鍵值段落。
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from .base import ParseError, ParseResult

# 全形/半形冒號皆視為鍵值分隔
_KV_PATTERN = re.compile(r"^\s*(.+?)\s*[:：]\s*(.*\S)?\s*$")


def _table_to_result(table) -> ParseResult | None:
    rows = table.rows
    if len(rows) < 2:
        return None
    header = [c.text.strip() for c in rows[0].cells]
    fields = [h for h in header if h != ""]
    if not fields:
        return None

    data_rows: list[dict[str, Any]] = []
    for r in rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        record: dict[str, Any] = {}
        has_value = False
        for idx, name in enumerate(header):
            if name == "":
                continue
            val = cells[idx] if idx < len(cells) else ""
            if val != "":
                has_value = True
            record[name] = val
        if has_value:
            data_rows.append(record)
    if not data_rows:
        return None
    return ParseResult(fields=fields, rows=data_rows)


def _paragraphs_to_result(doc) -> ParseResult | None:
    record: dict[str, Any] = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        m = _KV_PATTERN.match(text)
        if m:
            key = m.group(1).strip()
            val = (m.group(2) or "").strip()
            if key:
                record[key] = val
    if not record:
        return None
    return ParseResult(fields=list(record.keys()), rows=[record])


def parse(path: Path) -> ParseResult:
    from docx import Document

    doc = Document(str(path))

    # 先試表格,挑資料列最多的一個
    best_table: ParseResult | None = None
    for table in doc.tables:
        res = _table_to_result(table)
        if res and (best_table is None or len(res.rows) > len(best_table.rows)):
            best_table = res
    if best_table is not None:
        return best_table

    kv = _paragraphs_to_result(doc)
    if kv is not None:
        return kv

    raise ParseError(
        f"Word 檔「{path.name}」中找不到可辨識的表格或「欄位:值」內容。"
    )
